#!/usr/bin/env python3
"""Build the screenshot-based MES-lite operator guide source and optional DOCX."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    DOCX_IMPORT_ERROR: ModuleNotFoundError | None = None
except ModuleNotFoundError as error:
    DOCX_IMPORT_ERROR = error


ROOT = Path(__file__).resolve().parents[1]
DOCUMENT_FONT = "Arial Unicode MS"
EAST_ASIA_FONT = "Arial Unicode MS"
# compact_reference_guide 的横向现场截图覆盖：9.55in 固定表格，保留 120 DXA 左缩进。
TABLE_WIDTH_DXA = 13752
TABLE_INDENT_DXA = 120
SHOT_ROOT = ROOT / "docs/operations/user-guide/screenshots"
SOURCE_DIR = ROOT / "docs/operations/user-guide"
DOCX_DIR = ROOT / "output/docx"


CATALOG = json.loads((ROOT / "sop/manifest.json").read_text(encoding="utf-8"))
VERSION = json.loads((ROOT / "package.json").read_text(encoding="utf-8"))["version"]
CHAPTERS = [(chapter["title"], chapter["workflows"]) for chapter in CATALOG["chapters"]]
WORKFLOW_COUNT = sum(len(items) for _, items in CHAPTERS)
SCREENSHOT_BASELINES = sorted(
    {workflow["screenshot"]["baseline"] for _, items in CHAPTERS for workflow in items},
    key=lambda value: tuple(int(part) for part in value.split(".")),
)
SCREENSHOT_BASELINE_LABEL = f"v{SCREENSHOT_BASELINES[0]}-v{SCREENSHOT_BASELINES[-1]}"

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = DOCUMENT_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"表格列宽合计必须为 {TABLE_WIDTH_DXA} DXA：{widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                margin = tc_mar.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0


def add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.append(level_justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "271")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_paragraph(doc: Document, text: str, num_id: int, *, size: float = 11, compact: bool = False):
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    # 截图步骤页使用紧凑覆盖，其余列表保持 preset 的 4pt / 1.25。
    paragraph.paragraph_format.space_after = Pt(0 if compact else 4)
    paragraph.paragraph_format.line_spacing = 0.95 if compact else 1.25
    set_run_font(paragraph.add_run(text), size)
    return paragraph


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, 8, color="64748B")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, 8, color="64748B")


def apply_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DOCUMENT_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(30, 41, 59)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 28, "0F172A", 0, 3),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = DOCUMENT_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = DOCUMENT_FONT
    cap._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    cap.font.size = Pt(8)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(71, 85, 105)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(1)


def configure_section(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"MES-lite  全流程作业指导书  |  v{VERSION}")
    set_run_font(run, 8, True, "0F3B5F")
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def add_banner(doc: Document, text: str, fill: str = "E0F2FE", color: str = "075985") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 9.5, True, color)


def add_cover(doc: Document) -> None:
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MES-lite")
    set_run_font(run, 18, True, "0284C7")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("全流程作业指导书")
    set_run_font(run, 30, True, "0F172A")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("从主数据到生产、库存与销售闭环")
    set_run_font(run, 15, False, "475569")
    doc.add_paragraph("")
    add_banner(doc, "本指导书基于隔离的本地演示数据库和真实页面操作截图制作；不含生产业务数据。")
    doc.add_paragraph("")
    table = doc.add_table(rows=5, cols=2)
    set_repeat_table_header(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("交付版本", f"v{VERSION}"),
        ("截图基线", f"{SCREENSHOT_BASELINE_LABEL} 的 {WORKFLOW_COUNT} 张真实页面流程截图"),
        ("适用角色", "管理员、计划员、班组长、工艺、仓管、销售、人事、质检和系统维护人员"),
        ("数据范围", "业务页面使用隔离演示数据；恢复章节使用脱敏的生产候选演练证据"),
        ("编制日期", "2026-08-15"),
    ]
    for idx, (key, value) in enumerate(rows):
        table.cell(idx, 0).text = key
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), "E2E8F0")
        for run in table.cell(idx, 0).paragraphs[0].runs:
            set_run_font(run, 9.5, True, "0F172A")
        for run in table.cell(idx, 1).paragraphs[0].runs:
            set_run_font(run, 9.5)
    set_table_geometry(table, [2700, 11052])
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(18)
    run = p4.add_run("受控文档 · 操作前先确认权限、单据状态、物料、数量、单位与库位")
    set_run_font(run, 10, True, "B45309")


def add_front_matter(doc: Document, bullet_num_id: int) -> None:
    doc.add_page_break()
    front_heading = doc.add_heading("使用说明与边界", level=1)
    front_heading.paragraph_format.space_before = Pt(4)
    front_heading.paragraph_format.space_after = Pt(4)
    add_banner(doc, "先建立主数据，再创建业务单据；先核对草稿，再执行会改变库存或状态的确认动作。", "DCFCE7", "166534")
    bullets = [
        "绿色/成功提示只是页面提交成功；关键业务还必须在库存流水、订单详情或操作记录中复核。",
        "归档、永久删除、库存调整、数据修复、权限修改和 AI 密钥配置仅限授权人员。",
        "业务员工和登录账号是两套对象：员工用于单据执行，账号用于登录；绑定账号不会自动授予权限。",
        "生产订单统一状态：草稿 -> 已发布 -> 生产中 -> 已完成，或在允许阶段取消。草稿不可登记实绩。",
        "供应商批号、来料内部批次、生产投入、产出批次、客户发货、退货回流、质量判定与不合格处置已贯通；61 项细粒度资源、岗位任务、生产/来料/物流命令、设备事件、周期点检/维保、生产实绩设备/文件版本快照、数据范围和临时授权已落地；生产已建立固定提交、每日一致备份和失败通知，生产恢复候选已完成隔离应用启动、登录、业务与附件抽查；自动设备采集/OEE、Product 到 Material 人工确认和回填、异地副本、真实 Coolify 挂载切换与真实岗位审批仍待完成。",
        "旧 Product 兼容入口和旧生产领料/报工/QC/入库接口不作为本指导书主流程；当前主流程使用 Material、已发布 BOM 和班后实绩。",
    ]
    for item in bullets:
        add_list_paragraph(doc, item, bullet_num_id, size=9.5, compact=True)

    flow_heading = doc.add_heading("推荐业务顺序", level=2)
    flow_heading.paragraph_format.space_before = Pt(4)
    flow_heading.paragraph_format.space_after = Pt(3)
    flow = CATALOG["recommendedSequence"].replace("→", "->")
    add_banner(doc, flow)

    index_heading = doc.add_heading("章节索引", level=2)
    index_heading.paragraph_format.space_before = Pt(4)
    index_heading.paragraph_format.space_after = Pt(3)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["章节", "业务范围", "关键结果"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        set_cell_shading(table.cell(0, i), "0F3B5F")
        for run in table.cell(0, i).paragraphs[0].runs:
            set_run_font(run, 9, True, "FFFFFF")
    set_repeat_table_header(table.rows[0])
    chapter_rows = [
        ("1", "登录、工作台、MES/ERP 导航", "进入正确工作区"),
        ("2", "物料与 BOM", "统一主数据和已发布 BOM"),
        ("3-6", "生产订单、实绩、质检、批次谱系、派工、转移", "投入、产出和质量状态可回放"),
        ("7", "来料、库存、流水", "收发存可追溯"),
        ("8", "销售、发货、退货", "订单与库存闭环"),
        ("9-10", "文档、设备事件/点检、业务配置", "基础数据和点检事实可追溯"),
        ("11", "系统、工具、权限", "可配置、可审计、最小权限"),
        ("17", "Coolify 发布、备份和健康验收", "版本固定、数据可备份、故障可恢复"),
    ]
    for row in chapter_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs:
                set_run_font(run, 8.5)
    set_table_geometry(table, [1200, 6276, 6276])


def normalize_item(item):
    screenshot = item["screenshot"]
    return (
        screenshot["file"],
        item["title"],
        item["objective"],
        item["steps"],
        item["result"],
        screenshot["baseline"],
    )


def add_instruction_page(doc: Document, chapter: str, item, figure_number: int) -> None:
    filename, title, objective, steps, result, screenshot_baseline = normalize_item(item)
    image_path = SHOT_ROOT / f"v{screenshot_baseline}" / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    doc.add_page_break()
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(chapter)
    set_run_font(run, 10, True, "0284C7")
    step_heading = doc.add_heading(title, level=2)
    step_heading.paragraph_format.space_before = Pt(0)
    step_heading.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("目的：")
    set_run_font(r1, 9, True, "0F172A")
    r2 = p.add_run(objective)
    set_run_font(r2, 9)
    number_num_id = add_numbering_definition(doc, "decimal")
    for step in steps:
        add_list_paragraph(doc, step, number_num_id, size=8.4, compact=True)
    result_p = doc.add_paragraph()
    result_p.paragraph_format.space_after = Pt(3)
    rr1 = result_p.add_run("结果检查：")
    set_run_font(rr1, 8.8, True, "166534")
    rr2 = result_p.add_run(result)
    set_run_font(rr2, 8.8)
    picture_p = doc.add_paragraph()
    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_p.paragraph_format.space_after = Pt(0)
    picture_p.paragraph_format.keep_with_next = True
    picture = picture_p.add_run().add_picture(str(image_path), height=Inches(4.3))
    picture._inline.docPr.set("title", title)
    picture._inline.docPr.set("descr", f"MES-lite {title}页面操作截图")
    caption = doc.add_paragraph(style="Figure Caption")
    caption.paragraph_format.keep_together = True
    caption.add_run(f"图 {figure_number}  {title}（演示库截图：{filename}）")


def add_appendix(doc: Document, bullet_num_id: int) -> None:
    doc.add_page_break()
    appendix_a_heading = doc.add_heading("附录 A：上线前岗位检查表", level=1)
    appendix_a_heading.paragraph_format.space_before = Pt(4)
    appendix_a_heading.paragraph_format.space_after = Pt(4)
    checks = [
        ("计划/生产", "物料、BOM、计划数量和日期已复核；订单发布后再执行"),
        ("班组", "员工、工作中心、实际投入/产出、库位和班后日期真实"),
        ("仓储", "收货、发货、退货和调整均有来源单据；完成后核对流水"),
        ("销售", "客户、价格、交期、客户单号与发货占用一致"),
        ("文控", "文档版本、状态、适用物料和工作中心正确"),
        ("质检", "批次号、抽样结果和结论已核对；整批放行/冻结后复核库存状态与流水"),
        ("管理员", "账号已审核；质量和生产高风险命令使用独立权限；关键操作可审计；任务数量与岗位数据范围已复核"),
        ("运维", "操作前备份；恢复演练通过；密钥不进代码和截图"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(["岗位", "检查要点", "签字/日期"]):
        table.cell(0, i).text = h
        set_cell_shading(table.cell(0, i), "0F3B5F")
        for run in table.cell(0, i).paragraphs[0].runs:
            set_run_font(run, 9, True, "FFFFFF")
    for role, check in checks:
        row = table.add_row()
        cells = row.cells
        cells[0].text = role
        cells[1].text = check
        cells[2].text = ""
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                set_run_font(run, 9)
    set_table_geometry(table, [1900, 9252, 2600])
    appendix_b_heading = doc.add_heading("附录 B：暂不作为现行 SOP 的治理项", level=1)
    appendix_b_heading.paragraph_format.space_before = Pt(6)
    appendix_b_heading.paragraph_format.space_after = Pt(4)
    for item in CATALOG["governanceBoundaries"]:
        add_list_paragraph(doc, item, bullet_num_id, size=9.2, compact=True)
    add_banner(doc, "以上功能在完成数据模型、权限、回归测试和现场验证前，不应写入车间正式作业标准。", "FEF3C7", "92400E")


def build_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_section(doc)
    apply_styles(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet")
    doc.core_properties.title = "MES-lite 全流程作业指导书"
    doc.core_properties.subject = "MES-lite 业务操作与截图 SOP"
    doc.core_properties.author = "MES-lite 项目组"
    doc.core_properties.keywords = "MES-lite,MES,作业指导书,SOP"
    add_cover(doc)
    add_front_matter(doc, bullet_num_id)
    figure_number = 1
    for chapter, items in CHAPTERS:
        for item in items:
            add_instruction_page(doc, chapter, item, figure_number)
            figure_number += 1
    add_appendix(doc, bullet_num_id)
    doc.save(path)


def markdown_lines() -> Iterable[str]:
    yield "# MES-lite 全流程作业指导书"
    yield ""
    yield f"- 交付版本：v{VERSION}"
    yield f"- 截图基线：{SCREENSHOT_BASELINE_LABEL}（{WORKFLOW_COUNT} 张经验证流程截图）"
    yield "- 数据范围：业务流程使用隔离本地临时演示库；第 17 章使用真实 Coolify 运维配置和脱敏生产恢复候选演练证据，不展示生产业务明细、账号、密码或密钥"
    yield "- 编制日期：2026-08-14"
    yield ""
    yield f"> 重要：{CATALOG['important']}"
    yield ""
    yield "## 推荐业务顺序"
    yield ""
    yield CATALOG["recommendedSequence"]
    for chapter, items in CHAPTERS:
        yield ""
        yield f"## {chapter}"
        for item in items:
            filename, title, objective, steps, result, screenshot_baseline = normalize_item(item)
            yield ""
            yield f"### {title}"
            yield ""
            yield f"目的：{objective}"
            yield ""
            for idx, step in enumerate(steps, start=1):
                yield f"{idx}. {step}"
            yield ""
            yield f"结果检查：{result}"
            yield ""
            yield f"![{title}](screenshots/v{screenshot_baseline}/{filename})"
    yield ""
    yield "## 暂不作为现行 SOP 的治理项"
    yield ""
    for boundary in CATALOG["governanceBoundaries"]:
        yield f"- {boundary}"


def build_markdown(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(markdown_lines()) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=DOCX_DIR / f"MES-lite全流程作业指导书-v{VERSION}.docx")
    parser.add_argument("--markdown", type=Path, default=SOURCE_DIR / f"MES-lite全流程作业指导书-v{VERSION}.md")
    parser.add_argument("--markdown-only", action="store_true", help="只生成受 CI 管理的 Markdown 源文件")
    args = parser.parse_args()
    build_markdown(args.markdown)
    print(f"Created {args.markdown}")
    if not args.markdown_only:
        if DOCX_IMPORT_ERROR:
            raise RuntimeError("生成 DOCX 需要安装 python-docx；开发期请使用 --markdown-only") from DOCX_IMPORT_ERROR
        build_docx(args.docx)
        print(f"Created {args.docx}")


if __name__ == "__main__":
    main()
