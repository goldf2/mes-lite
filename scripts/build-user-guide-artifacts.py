#!/usr/bin/env python3
"""Build the versioned MES-lite screenshot SOP as stable page-image DOCX/PDF.

Each workflow is rendered once to a self-contained A4 page. The same page images
are embedded in DOCX and PDF so Chinese text and layout do not depend on fonts
installed on the reader's computer.
"""

from __future__ import annotations

import re
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


VERSION = "v0.1.362"
EXPECTED_WORKFLOWS = 138
PAGE_W = 1240
PAGE_H = 1754
MARGIN = 72
BLUE = "#2563EB"
DARK = "#172033"
MUTED = "#667085"
LIGHT_BLUE = "#EFF6FF"
LIGHT_GRAY = "#F6F8FB"
GREEN = "#047857"
LIGHT_GREEN = "#ECFDF3"
WHITE = "#FFFFFF"
FONT_PATHS = (
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/PingFang.ttc"),
    Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
)


@dataclass
class Workflow:
    chapter: str
    title: str
    lines: list[str]
    image_path: Path
    image_alt: str


def font_path() -> Path:
    for candidate in FONT_PATHS:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("找不到可用于生成指导书的中文字体")


FONT_PATH = font_path()


def font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size=size)


def text_height(draw: ImageDraw.ImageDraw, value: str, face: ImageFont.FreeTypeFont) -> int:
    box = draw.textbbox((0, 0), value or "国", font=face)
    return box[3] - box[1]


def text_width(draw: ImageDraw.ImageDraw, value: str, face: ImageFont.FreeTypeFont) -> int:
    return int(draw.textlength(value, font=face))


def wrap_text(draw: ImageDraw.ImageDraw, value: str, face: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    value = value.strip()
    if not value:
        return []
    rows: list[str] = []
    current = ""
    for char in value:
        candidate = current + char
        if current and text_width(draw, candidate, face) > max_width:
            rows.append(current.rstrip())
            current = char.lstrip()
        else:
            current = candidate
    if current:
        rows.append(current.rstrip())
    return rows


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    value: str,
    face: ImageFont.FreeTypeFont,
    fill: str,
    max_width: int,
    line_gap: int = 10,
) -> int:
    x, y = xy
    line_h = text_height(draw, "国Ag", face) + line_gap
    for row in wrap_text(draw, value, face, max_width):
        draw.text((x, y), row, font=face, fill=fill)
        y += line_h
    return y


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str | None = None) -> None:
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=2 if outline else 1)


def new_page() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    page = Image.new("RGB", (PAGE_W, PAGE_H), WHITE)
    return page, ImageDraw.Draw(page)


def draw_header(draw: ImageDraw.ImageDraw, section: str, page_no: int, total_pages: int) -> None:
    draw.text((MARGIN, 42), "MES-lite  ·  全流程作业指导书", font=font(22), fill=DARK)
    version_text = f"{VERSION}  ·  第 {page_no}/{total_pages} 页"
    draw.text((PAGE_W - MARGIN - text_width(draw, version_text, font(20)), 45), version_text, font=font(20), fill=MUTED)
    draw.line((MARGIN, 86, PAGE_W - MARGIN, 86), fill="#D9E1EC", width=2)
    draw.text((MARGIN, 112), section, font=font(22), fill=BLUE)


def draw_footer(draw: ImageDraw.ImageDraw) -> None:
    draw.line((MARGIN, PAGE_H - 58, PAGE_W - MARGIN, PAGE_H - 58), fill="#E4E7EC", width=2)
    footer = "隔离本地临时演示库  ·  内部受控文件  ·  使用前请核对系统版本"
    draw.text((MARGIN, PAGE_H - 43), footer, font=font(17), fill=MUTED)


def parse_source(source_path: Path) -> tuple[list[str], str, list[Workflow], list[str]]:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    metadata: list[str] = []
    important = ""
    workflows: list[Workflow] = []
    boundaries: list[str] = []
    chapter = ""
    title = ""
    content: list[str] = []
    mode = "intro"

    def flush() -> None:
        nonlocal title, content
        if not title:
            return
        image_match = next(
            (
                re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", line.strip())
                for line in content
                if line.strip().startswith("![")
            ),
            None,
        )
        if not image_match:
            raise ValueError(f"流程缺少截图：{title}")
        image_path = (source_path.parent / image_match.group(2)).resolve()
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        body = [line for line in content if not line.strip().startswith("![")]
        workflows.append(Workflow(chapter, title, body, image_path, image_match.group(1)))
        title = ""
        content = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("## "):
            flush()
            chapter = line[3:].strip()
            mode = "boundaries" if chapter == "暂不作为现行 SOP 的治理项" else "other"
            continue
        if line.startswith("### "):
            flush()
            title = line[4:].strip()
            mode = "workflow"
            continue
        if title:
            content.append(line)
        elif mode == "intro":
            if line.startswith("- "):
                metadata.append(line[2:])
            elif line.startswith("> "):
                important = line[2:]
        elif mode == "boundaries" and line.startswith("- "):
            boundaries.append(line[2:])
    flush()
    return metadata, important, workflows, boundaries


def render_cover(metadata: list[str], important: str, workflow_count: int, total_pages: int) -> Image.Image:
    page, draw = new_page()
    draw.rounded_rectangle((0, 0, PAGE_W, 350), radius=0, fill="#F4F7FF")
    draw.text((MARGIN, 74), "MES · MRP-lite · ERP-lite", font=font(25), fill=BLUE)
    draw.text((MARGIN, 146), "MES-lite", font=font(66), fill=DARK)
    draw.text((MARGIN, 235), "全流程作业指导书", font=font(58), fill=DARK)
    draw.text((MARGIN, 382), f"{VERSION}  ·  {workflow_count} 个实操流程  ·  截图版 SOP", font=font(28), fill=MUTED)

    y = 470
    rounded_box(draw, (MARGIN, y, PAGE_W - MARGIN, y + 310), LIGHT_BLUE)
    y += 34
    draw.text((MARGIN + 32, y), "文档基线", font=font(27), fill=BLUE)
    y += 58
    for item in metadata:
        key, sep, value = item.partition("：")
        label = f"{key}：" if sep else ""
        draw.text((MARGIN + 32, y), label, font=font(23), fill=BLUE)
        draw_wrapped(draw, (MARGIN + 210, y), value if sep else item, font(23), DARK, PAGE_W - 2 * MARGIN - 250, 8)
        y += 53

    y = 840
    rounded_box(draw, (MARGIN, y, PAGE_W - MARGIN, y + 330), LIGHT_GRAY)
    draw.text((MARGIN + 32, y + 30), "适用边界", font=font(28), fill=DARK)
    draw_wrapped(draw, (MARGIN + 32, y + 90), important, font(25), MUTED, PAGE_W - 2 * MARGIN - 64, 14)

    y = 1245
    rounded_box(draw, (MARGIN, y, PAGE_W - MARGIN, y + 230), "#172033")
    draw.text((MARGIN + 32, y + 30), "使用方法", font=font(27), fill="#93C5FD")
    draw_wrapped(
        draw,
        (MARGIN + 32, y + 87),
        "每个流程独立成页。操作者按步骤执行，并用同页的“结果检查”和实机截图确认操作结果。上线环境须使用授权账号和真实业务单据复核。",
        font(25),
        WHITE,
        PAGE_W - 2 * MARGIN - 64,
        12,
    )
    draw_footer(draw)
    return page


def render_workflow(workflow: Workflow, index: int, total: int, page_no: int, total_pages: int) -> Image.Image:
    page, draw = new_page()
    draw_header(draw, f"{workflow.chapter}  ·  流程 {index:02d}/{total:02d}", page_no, total_pages)
    y = 160
    y = draw_wrapped(draw, (MARGIN, y), workflow.title, font(42), DARK, PAGE_W - 2 * MARGIN, 12) + 12

    purpose = ""
    result = ""
    steps: list[str] = []
    for raw in workflow.lines:
        line = raw.strip()
        if line.startswith("目的："):
            purpose = line.removeprefix("目的：")
        elif line.startswith("结果检查："):
            result = line.removeprefix("结果检查：")
        elif re.match(r"^\d+\.\s+", line):
            steps.append(line)

    if purpose:
        purpose_lines = wrap_text(draw, purpose, font(25), PAGE_W - 2 * MARGIN - 184)
        height = max(80, 36 + len(purpose_lines) * 40)
        rounded_box(draw, (MARGIN, y, PAGE_W - MARGIN, y + height), LIGHT_BLUE)
        draw.text((MARGIN + 24, y + 24), "目的", font=font(25), fill=BLUE)
        draw_wrapped(draw, (MARGIN + 128, y + 22), purpose, font(25), DARK, PAGE_W - 2 * MARGIN - 184, 9)
        y += height + 20

    draw.text((MARGIN, y), "操作步骤", font=font(27), fill=DARK)
    y += 48
    for step in steps:
        number, value = step.split(".", 1)
        circle_y = y + 14
        draw.ellipse((MARGIN, circle_y - 15, MARGIN + 38, circle_y + 23), fill=BLUE)
        no_face = font(20)
        no_w = text_width(draw, number, no_face)
        draw.text((MARGIN + (38 - no_w) / 2, circle_y - 12), number, font=no_face, fill=WHITE)
        next_y = draw_wrapped(draw, (MARGIN + 58, y), value.strip(), font(25), DARK, PAGE_W - 2 * MARGIN - 58, 9)
        y = max(y + 44, next_y) + 8

    if result:
        result_lines = wrap_text(draw, result, font(24), PAGE_W - 2 * MARGIN - 206)
        height = max(76, 32 + len(result_lines) * 39)
        rounded_box(draw, (MARGIN, y, PAGE_W - MARGIN, y + height), LIGHT_GREEN)
        draw.text((MARGIN + 24, y + 22), "结果检查", font=font(24), fill=GREEN)
        draw_wrapped(draw, (MARGIN + 158, y + 20), result, font(24), DARK, PAGE_W - 2 * MARGIN - 206, 9)
        y += height + 22

    caption_h = 42
    max_image_h = PAGE_H - 82 - caption_h - y
    with Image.open(workflow.image_path) as source:
        screenshot = source.convert("RGB")
    max_image_w = PAGE_W - 2 * MARGIN
    scale = min(max_image_w / screenshot.width, max_image_h / screenshot.height)
    if scale <= 0:
        raise ValueError(f"页面内容超高，无法放置截图：{workflow.title}")
    width = max(1, int(screenshot.width * scale))
    height = max(1, int(screenshot.height * scale))
    screenshot = screenshot.resize((width, height), Image.Resampling.LANCZOS)
    x = (PAGE_W - width) // 2
    draw.rounded_rectangle((x - 3, y - 3, x + width + 3, y + height + 3), radius=9, fill="#CBD5E1")
    page.paste(screenshot, (x, y))
    caption = f"图 {index:02d}  {workflow.image_alt}"
    caption_x = (PAGE_W - text_width(draw, caption, font(18))) // 2
    draw.text((caption_x, y + height + 13), caption, font=font(18), fill=MUTED)
    draw_footer(draw)
    return page


def render_boundaries(boundaries: list[str], page_no: int, total_pages: int) -> Image.Image:
    page, draw = new_page()
    draw_header(draw, "治理边界", page_no, total_pages)
    y = 176
    draw.text((MARGIN, y), "暂不作为现行 SOP 的治理项", font=font(43), fill=DARK)
    y += 92
    rounded_box(draw, (MARGIN, y, PAGE_W - MARGIN, y + 118), LIGHT_BLUE)
    draw_wrapped(
        draw,
        (MARGIN + 28, y + 23),
        "以下能力仍处于治理路线中，不能在当前版本中宣称已经形成生产闭环。",
        font(26),
        DARK,
        PAGE_W - 2 * MARGIN - 56,
        10,
    )
    y += 164
    for index, item in enumerate(boundaries, start=1):
        rows = wrap_text(draw, item, font(27), PAGE_W - 2 * MARGIN - 94)
        height = max(88, 38 + len(rows) * 43)
        rounded_box(draw, (MARGIN, y, PAGE_W - MARGIN, y + height), LIGHT_GRAY)
        draw.text((MARGIN + 24, y + 23), f"{index:02d}", font=font(25), fill=BLUE)
        draw_wrapped(draw, (MARGIN + 86, y + 21), item, font(27), DARK, PAGE_W - 2 * MARGIN - 112, 10)
        y += height + 20
    draw_footer(draw)
    return page


def save_docx(page_paths: list[Path], output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.1)
    section.bottom_margin = Cm(0.1)
    section.left_margin = Cm(0.2)
    section.right_margin = Cm(0.2)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = 0
    normal.paragraph_format.space_after = 0
    normal.paragraph_format.line_spacing = 1
    for index, page_path in enumerate(page_paths):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        # Leave enough room for Writer/Word's unavoidable paragraph end mark;
        # otherwise a full-height page image creates one extra blank page.
        picture = paragraph.add_run().add_picture(str(page_path), width=Cm(20.2))
        picture._inline.docPr.set("title", f"MES-lite 作业指导书第 {index + 1} 页")
        picture._inline.docPr.set("descr", f"MES-lite 全流程作业指导书第 {index + 1} 页完整页面")
        if index < len(page_paths) - 1:
            doc.add_page_break()
    props = doc.core_properties
    props.title = "MES-lite 全流程作业指导书"
    props.subject = "MES-lite 业务操作、批次追溯与截图 SOP"
    props.author = "MES-lite 项目组"
    props.keywords = "MES-lite,MES,作业指导书,SOP,批次追溯,退货质检"
    props.comments = "versioned A4 pages generated from Markdown and real local screenshots"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def save_pdf(page_paths: list[Path], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(output_path), pagesize=A4, pageCompression=1)
    page_w, page_h = A4
    for page_path in page_paths:
        pdf.drawImage(str(page_path), 0, 0, width=page_w, height=page_h, preserveAspectRatio=True, mask="auto")
        pdf.showPage()
    pdf.setTitle("MES-lite 全流程作业指导书")
    pdf.setAuthor("MES-lite 项目组")
    pdf.setSubject("MES-lite 业务操作、批次追溯与截图 SOP")
    pdf.save()


def main() -> int:
    if len(sys.argv) != 4:
        print("usage: build-user-guide-artifacts.py SOURCE.md OUTPUT.docx OUTPUT.pdf", file=sys.stderr)
        return 2
    source_path = Path(sys.argv[1]).resolve()
    docx_path = Path(sys.argv[2]).resolve()
    pdf_path = Path(sys.argv[3]).resolve()
    metadata, important, workflows, boundaries = parse_source(source_path)
    if len(workflows) != EXPECTED_WORKFLOWS:
        raise ValueError(f"预期 {EXPECTED_WORKFLOWS} 个截图流程，实际 {len(workflows)}")

    total_pages = len(workflows) + 2
    with tempfile.TemporaryDirectory(prefix="mes-lite-guide-pages-") as temp_dir:
        page_paths: list[Path] = []
        pages = [render_cover(metadata, important, len(workflows), total_pages)]
        pages.extend(
            render_workflow(workflow, index, len(workflows), index + 1, total_pages)
            for index, workflow in enumerate(workflows, start=1)
        )
        pages.append(render_boundaries(boundaries, total_pages, total_pages))
        for index, page in enumerate(pages, start=1):
            page_path = Path(temp_dir) / f"page-{index:03d}.jpg"
            page.save(page_path, "JPEG", quality=91, optimize=True, subsampling=0)
            page_paths.append(page_path)
        save_docx(page_paths, docx_path)
        save_pdf(page_paths, pdf_path)

    print(f"DOCX generated: {docx_path} ({len(workflows)} workflows, {total_pages} pages)")
    print(f"PDF generated:  {pdf_path} ({len(workflows)} workflows, {total_pages} pages)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
